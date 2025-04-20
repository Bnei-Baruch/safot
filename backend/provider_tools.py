"""
This file is not part of the main application.
It's a command-line tool for translating documents directly from the terminal.
Usage: python provider_tools.py -i input.docx -o output.docx -t Hebrew
"""

import os
from dotenv import load_dotenv
from provider_libs import TranslationProvider
from docx import Document
from io import BytesIO
import argparse


load_dotenv()
parser = argparse.ArgumentParser(description="Providers tool")
DEFAULT_PROMPT = (
    "Translate the following text into %(target_language)s, ensuring the translation is faithful to the original meaning "
    "and context of the teachings of Kabbalah. This text is an excerpt from the book titled 'Thank You to Women', "
    "which explores the spiritual role and significance of women in the wisdom of Kabbalah. "
    "The book is intended for an audience deeply interested in Kabbalistic teachings and spiritual development, "
    "seeking to understand the unique spiritual contributions and roles of women according to authentic Kabbalistic sources.\n\n"
    "The translation should:\n"
    "1. Accurately reflect the unique terminology and style of Kabbalistic wisdom, as taught by authentic Kabbalists like Baal HaSulam and Rabash.\n"
    "2. Maintain a respectful and precise tone, suitable for readers invested in the spiritual significance of the role of women in Kabbalah.\n"
    "3. Avoid rephrasing or altering the intent of the original text, while ensuring linguistic fluency and readability in the target language.\n"
    "4. Preserve any cultural or conceptual nuances significant to the original message.\n\n"
    "Separate each paragraph using the `|||` separator exactly as it appears in the original text.\n\n"
    "Begin translating:"
)

parser.add_argument("-i", "--input", type=str, help="Docx to translate")
parser.add_argument("-o", "--output", type=str, help="Translated docx")
parser.add_argument("-m", "--model", type=str,
                    default="gpt-3.5-turbo", help="Which model to use")
parser.add_argument("-t", "--target_language", type=str,
                    default="Hebrew", help="Target language for translation")
parser.add_argument("-p", "--prompt", type=str,
                    default=DEFAULT_PROMPT, help="Translation prompt")
parser.add_argument("-a", "--apply", type=bool, default=False,
                    help="If false will not really call the API")

args = parser.parse_args()


def generate_docx(paragraphs, output_filename):
    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    doc.save(output_filename)


def main():
    api_key = os.environ.get("OPENAI_API_KEY")
    paragraphs = []
    # Read the input file
    with open(args.input, 'rb') as file:
        content = file.read()
        document = Document(BytesIO(content))
        for p in document.paragraphs:
            paragraphs.append(p.text)

    print(f"Translating {len(paragraphs)} paragraphs.")
    translated_paragraphs = []
    if args.apply:
        # Create a TranslationProvider object
        translator = TranslationProvider(
            api_key=api_key,
            model=args.model,
            prompt=args.prompt,
            target_language=args.target_language,
            paragraphs=paragraphs
        )

        # Run the translation process
        print("Starting translation...")
        translated_paragraphs = translator.translate_text(output_ratio=1.2)
    else:
        translated_paragraphs = [f"Tr {p}" for p in paragraphs]

    # Display the translated paragraphs
    print("\nTranslated Paragraphs:")
    for i, paragraph in enumerate(translated_paragraphs, 1):
        print(f"Paragraph {i}: {paragraph}")

    generate_docx(translated_paragraphs, args.output)


if __name__ == "__main__":
    main()
