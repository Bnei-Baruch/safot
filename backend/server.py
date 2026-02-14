from datetime import datetime, timezone
from db import db
from docx import Document
from dotenv import load_dotenv
from keycloak import KeycloakOpenID
import logging
import os
import traceback

from peewee import (
    SQL,
    JOIN,
    fn,
)
from peewee_migrate import Router
from playhouse.shortcuts import model_to_dict
from starlette.status import HTTP_401_UNAUTHORIZED

from fastapi import (
    Depends,
    FastAPI,
    File,
    Form,
    HTTPException,
    Query,
    Request,
    UploadFile,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse

from services.translation_service import get_provider_lock
from services.provider_factory import create_translation_provider
from services.openai_provider import OPENAI_MODELS, PROVIDER_NAME as OPENAI_NAME, PROVIDER_LABEL as OPENAI_LABEL
from services.claude_provider import CLAUDE_MODELS, PROVIDER_NAME as CLAUDE_NAME, PROVIDER_LABEL as CLAUDE_LABEL
from services.segment_service import get_paragraphs_from_file, get_latest_segments, store_segments
from services.source_service import (
    create_or_update_sources,
    get_sources,
)
from services.dictionary import (
  get_dictionaries,
  get_rules,
)
from services.prompt import (
    build_prompt_from_dictionary,
    get_task_prompt,
    LANGUAGES,
    SEGMENTS_SUFFIX,
    RULE_TYPE_TEXT,
    RULE_TYPE_SEGMENTS_SUFFIX,
)
from services.utils import (
    apply_dict,
    epoch_microseconds,
    microseconds,
    to_datetime,
)

from models import (
    Dictionaries,
    ParagraphsTranslateRequest,
    PromptRequest,
    Provider,
    Rules,
    Segments,
    SegmentsOrigins,
    Sources,
    SourcesOrigins,
    TranslationServiceOptions,
)

def configure_logging():
    """Configure logging for the entire application"""
    logging.basicConfig(
        level=logging.DEBUG,
        format='%(asctime)s - %(name)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s'
    )
    # Configure peewee logger
    peewee_logger = logging.getLogger('peewee')
    peewee_logger.setLevel(logging.DEBUG)

# Configure logging at startup
configure_logging()

logger = logging.getLogger(__name__)

def required(param, error_message: str):
    """Validate that a required parameter is present, raise 400 if not."""
    if not param:
        raise HTTPException(status_code=400, detail=error_message)

load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Keycloak
logger.info('Initializing keycloak with URL: %s, Client ID: %s, Realm: %s', 
            os.getenv('KEYCLOAK_SERVER_URL'),
            os.getenv('KEYCLOAK_CLIENT_ID'), 
            os.getenv('KEYCLOAK_REALM_NAME'))
keycloak_openid = KeycloakOpenID(
    server_url=os.getenv('KEYCLOAK_SERVER_URL'),
    client_id=os.getenv('KEYCLOAK_CLIENT_ID'),
    realm_name=os.getenv('KEYCLOAK_REALM_NAME'),
)

# User info middlware

async def get_user_info(request: Request):
    # Extract token from Authorization header
    auth_header = request.headers.get('Authorization')
    if not auth_header or not auth_header.startswith('Bearer '):
        logger.error('Missing authorization header')
        raise HTTPException(status_code=HTTP_401_UNAUTHORIZED,
                            detail='Missing or invalid token')

    token = auth_header[len('Bearer '):].strip()
    try:
        # Validate token and get user info
        user_info = keycloak_openid.userinfo(token)
        return user_info
    except Exception as e:
        logger.error('Invalid or expired token: %s', e)
        raise HTTPException(status_code=HTTP_401_UNAUTHORIZED,
                            detail='Invalid or expired token')


@app.on_event('startup')
def startup():
    if db.is_closed():
        db.connect()
    
    router = Router(db, migrate_dir='migrations')
    router.run()

    logger.info('Database connected and migrations applied')

@app.on_event('shutdown')
def shutdown():
    if not db.is_closed():
        db.close()
    logger.info('Database connection closed')

####### SOURCES
@app.post('/sources', response_model=list[dict])
async def sources_handler(request: Request, metadata: bool = Query(False), user_info: dict = Depends(get_user_info)):
    """Create/update or fetch sources
    - If body has 'source_ids' key: fetch sources by IDs (empty array = all sources)
    - Otherwise: create/update sources
    """
    data = await request.json()

    if 'source_ids' in data:
        # Fetch sources
        source_ids = data.get("source_ids", [])
        if not source_ids or len(source_ids) == 0:
            return get_sources(metadata)
        return get_sources(metadata, source_ids)
    else:
        # Create/update sources
        sources = data if isinstance(data, list) else [data]
        return create_or_update_sources(sources, user_info['preferred_username'])

@app.post('/sources/relations', response_model=list[dict])
def get_source_relations(request: dict, user_info: dict = Depends(get_user_info)):
    """Get source relations for a list of source IDs

    Request body: { "source_ids": [1, 2, 3, ...] }
    Returns: [{"origin_source_id": 1, "translated_source_id": 2}, ...]
    """
    source_ids = request.get("source_ids", [])
    if not source_ids:
        return []

    relations = list(SourcesOrigins.select(
        SourcesOrigins.origin_source_id,
        SourcesOrigins.translated_source_id
    ).where(
        (SourcesOrigins.translated_source_id.in_(source_ids)) |
        (SourcesOrigins.origin_source_id.in_(source_ids))
    ).dicts())

    return relations

@app.post('/sources/origins', response_model=list[dict])
async def create_source_origin_links(request: Request, user_info: dict = Depends(get_user_info)):
    """Create source origin relations

    Request body: { "relations": [{"origin_source_id": 1, "translated_source_id": 2}, ...] }
    Returns: [{"origin_source_id": 1, "translated_source_id": 2}, ...]
    """
    try:
        data = await request.json()
        relations = data.get("relations", [])

        if not isinstance(relations, list):
            raise HTTPException(status_code=400, detail="relations must be a list")

        created_links = []

        for rel in relations:
            origin_source_id = rel.get("origin_source_id")
            translated_source_id = rel.get("translated_source_id")

            if origin_source_id is None or translated_source_id is None:
                raise HTTPException(status_code=400, detail="Each relation must have origin_source_id and translated_source_id")

            # Get next sequence value for id
            cursor = db.execute_sql("SELECT nextval('sources_origins_id_seq')")
            link_id = cursor.fetchone()[0]

            # Create the link
            sources_origins = SourcesOrigins.create(
                id=link_id,
                origin_source_id=origin_source_id,
                translated_source_id=translated_source_id,
            )
            created_links.append({
                "origin_source_id": sources_origins.origin_source_id,
                "translated_source_id": sources_origins.translated_source_id
            })

        logger.info(f"Created {len(created_links)} source origin link(s)")
        return created_links
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Error creating source origin links: %s", e)
        raise HTTPException(status_code=500, detail=f"Failed to create source origin links: {str(e)}")

@app.delete('/sources/{source_id}', response_model=list)
def delete_source(source_id: int, _: dict = Depends(get_user_info)):
    try:
        source = Sources.get(Sources.id == source_id)
    except Exception as e:
        logger.error(f"Deletion - Source not found: {e}")
        raise HTTPException(status_code=404, detail='Source not found')

    # Check if this source has any translations (other sources use it as origin)
    translations = list(SourcesOrigins.select().where(
        SourcesOrigins.origin_source_id == source_id
    ))
    if translations:
        raise HTTPException(
            status_code=400,
            detail=f'Cannot delete source: {len(translations)} translation(s) depend on it'
        )

    # Get all segments for this source
    segment_ids = [seg.id for seg in Segments.select(Segments.id).where(
        Segments.source_id == source_id
    )]

    # Check if any segments have translations (other segments use them as origin)
    if segment_ids:
        segment_translations = list(SegmentsOrigins.select().where(
            SegmentsOrigins.origin_segment_id.in_(segment_ids)
        ))
        if segment_translations:
            raise HTTPException(
                status_code=400,
                detail=f'Cannot delete source: {len(segment_translations)} segment translation(s) depend on it'
            )

    # Delete all segment relations for these segments
    if segment_ids:
        SegmentsOrigins.delete().where(
            (SegmentsOrigins.origin_segment_id.in_(segment_ids)) |
            (SegmentsOrigins.translated_segment_id.in_(segment_ids))
        ).execute()

    # Delete all source relations for this source
    SourcesOrigins.delete().where(
        (SourcesOrigins.origin_source_id == source_id) |
        (SourcesOrigins.translated_source_id == source_id)
    ).execute()

    # Delete all segments for this source
    Segments.delete().where(Segments.source_id == source_id).execute()

    # Delete the source itself
    Sources.delete().where(Sources.id == source_id).execute()

    return [source_id]

####### SEGMENTS 
# TODO: Refactor segments to have _epoch fields and created_/modified_ fields..
@app.post('/segments', response_model=list[dict])
async def segments_handler(request: Request, user_info: dict = Depends(get_user_info)):
    """Fetch or save segments
    - If body has 'source_ids' key: fetch segments from those sources (empty array = all segments)
    - Otherwise: save segments
    """
    try:
        data = await request.json()

        if 'source_ids' in data:
            # Fetch segments
            source_ids = data.get("source_ids", [])
            # Empty list means all segments
            return get_latest_segments(source_ids if len(source_ids) > 0 else None)
        else:
            # Save segments
            segments = data.get("segments", [])

            if not isinstance(segments, list):
                raise HTTPException(status_code=400, detail="Invalid request format - segments must be a list")

            # Add username and timestamp to each segment
            now = datetime.now(timezone.utc)
            for segment in segments:
                segment["username"] = user_info["preferred_username"]
                segment["timestamp"] = now

            saved_segments = store_segments(segments)
            return saved_segments

    except Exception as e:
        logger.error("Error in /segments: %s", e)
        raise HTTPException(status_code=500, detail=f"Failed to process segments: {str(e)}")

@app.post('/segments/origins', response_model=list[dict])
async def create_segment_origin_links(request: Request, user_info: dict = Depends(get_user_info)):
    """Create segment origin relations

    Request body: { "relations": [{"origin_segment_id": 1, "origin_segment_timestamp": "2024-01-01T00:00:00Z", "translated_segment_id": 2, "translated_segment_timestamp": "2024-01-01T00:00:00Z"}, ...] }
    Timestamps can be ISO format strings or epoch microseconds integers.
    Returns: [{"origin_segment_id": 1, "origin_segment_timestamp": "2024-01-01T00:00:00Z", "translated_segment_id": 2, "translated_segment_timestamp": "2024-01-01T00:00:00Z"}, ...]
    """
    try:
        data = await request.json()
        relations = data.get("relations", [])

        print(f"DEBUG: Exception type is {type(Exception)} and value is {Exception}")
        if not isinstance(relations, list):
            raise HTTPException(status_code=400, detail="relations must be a list")

        created_links = []

        for rel in relations:
            origin_segment_id = rel.get("origin_segment_id")
            origin_segment_timestamp = rel.get("origin_segment_timestamp")
            translated_segment_id = rel.get("translated_segment_id")
            translated_segment_timestamp = rel.get("translated_segment_timestamp")

            if origin_segment_id is None or origin_segment_timestamp is None or translated_segment_id is None or translated_segment_timestamp is None:
                raise HTTPException(status_code=400, detail="Each relation must have origin_segment_id, origin_segment_timestamp, translated_segment_id, and translated_segment_timestamp")

            # Convert timestamps to datetime objects (handles both ISO strings and epoch microseconds)
            origin_segment_timestamp = to_datetime(origin_segment_timestamp)
            translated_segment_timestamp = to_datetime(translated_segment_timestamp)

            # Get next sequence value for id
            cursor = db.execute_sql("SELECT nextval('segments_origins_id_seq')")
            link_id = cursor.fetchone()[0]

            # Create the link
            segments_origins = SegmentsOrigins.create(
                id=link_id,
                origin_segment_id=origin_segment_id,
                origin_segment_timestamp=origin_segment_timestamp,
                translated_segment_id=translated_segment_id,
                translated_segment_timestamp=translated_segment_timestamp,
            )
            created_links.append({
                "origin_segment_id": segments_origins.origin_segment_id,
                "origin_segment_timestamp": segments_origins.origin_segment_timestamp.isoformat(),
                "translated_segment_id": segments_origins.translated_segment_id,
                "translated_segment_timestamp": segments_origins.translated_segment_timestamp.isoformat()
            })

        logger.info(f"Created {len(created_links)} segment origin link(s)")
        return created_links
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Error creating segment origin links: %s", e)
        raise HTTPException(status_code=500, detail=f"Failed to create segment origin links: {str(e)}")

####### TRANSLATION
@app.get("/providers", response_model=list)
def get_providers_handler():
    """
    Get list of available translation providers and their models.
    Returns provider metadata for frontend UI.
    """
    try:
        providers = [
            {
                "value": OPENAI_NAME,
                "label": OPENAI_LABEL,
                "models": OPENAI_MODELS
            },
            {
                "value": CLAUDE_NAME,
                "label": CLAUDE_LABEL,
                "models": CLAUDE_MODELS
            }
        ]
        return providers
    except Exception as e:
        logger.error("Error getting providers: %s", e)
        raise HTTPException(status_code=500, detail=f"Failed to get providers: {str(e)}")

@app.post("/translate", response_model=dict)
def translate_paragraphs_handler(
    request: ParagraphsTranslateRequest,
    user_info: dict = Depends(get_user_info)
):
    try:
        start_time = datetime.now(timezone.utc)

        required(request.paragraphs, "No paragraphs provided.")
        required(request.original_language, "Missing original_language in request.")
        required(request.translate_language, "Missing translate_language in request.")

        if request.additional_sources_languages and (not request.additional_sources_texts or len(request.additional_sources_texts) != len(request.additional_sources_languages)):
            raise HTTPException(status_code=400, detail="len(additional_sources_texts) should match len(additional_sources_languages).")

        # Determine provider and model (defaults for backward compatibility)
        provider = request.provider if request.provider else Provider.OPENAI

        # Set default model based on provider if not specified
        if request.model:
            model = request.model
        else:
            # Default models per provider
            if provider == Provider.CLAUDE:
                model = "claude-sonnet-4-5-20250929"
            else:
                model = "gpt-4o"

        options = TranslationServiceOptions(
            provider=provider,
            model=model,
            temperature=0.2,
            tpm_limit=30000
        )

        # Create provider instance using factory
        translation_service = create_translation_provider(provider, options)

        # Acquire provider lock to prevent concurrent translations that would exceed TPM limit
        provider_lock = get_provider_lock(options.provider.value)

        logger.info(f"User {user_info['preferred_username']} waiting for {options.provider.value} translation lock...")
        with provider_lock:
            logger.info(f"User {user_info['preferred_username']} acquired {options.provider.value} translation lock")
            result = translation_service.translate_paragraphs(
                original_language=request.original_language,
                paragraphs=request.paragraphs,
                additional_sources_languages=request.additional_sources_languages,
                additional_sources_texts=request.additional_sources_texts,
                translate_language=request.translate_language,
                task_prompt=request.task_prompt,
            )
            logger.info(f"User {user_info['preferred_username']} released {options.provider.value} translation lock")

        # Convert references_by_language dict to additional_sources_paragraphs list for backward compatibility
        # The order must match additional_sources_languages
        additional_sources_paragraphs = []
        for lang_code in request.additional_sources_languages:
            lang_name = LANGUAGES.get(lang_code, lang_code)
            refs = result["references_by_language"].get(lang_name, [])
            additional_sources_paragraphs.append(refs)

        end_time = datetime.now(timezone.utc)
        total_duration = (end_time - start_time).total_seconds()
        logger.info("Total translation time: %.2f seconds for %d paragraphs", total_duration, len(request.paragraphs))

        return {
            "translated_paragraphs": result["translated_paragraphs"],
            "additional_sources_paragraphs": additional_sources_paragraphs,
            "remaining_additional_sources_texts": result["remaining_additional_sources_texts"],
            "properties": result["properties"],
            "total_segments_translated": len(result["translated_paragraphs"]),
            "translation_time_seconds": total_duration
        }

    except Exception as e:
        logger.error("Error in translation handler: %s", e)
        logger.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Translation failed: {str(e)}")
    
####### IMPORT/EXPORT
# TODO: Consider moving this to a frontend library.
@app.post('/docx2text')
def extract_paragraphs_handler(files: list[UploadFile] = File(...)):
    try:
        results = []
        for file in files:
            if not file.filename.lower().endswith('.docx'):
                raise HTTPException(status_code=400, detail=f"Only .docx files are supported. Invalid file: {file.filename}")

            paragraphs = get_paragraphs_from_file(file)
            results.append(paragraphs)

        return results
    except HTTPException as e:
        raise e
    except Exception as e:
        logger.error("Error in /docx2text: %s", e)
        raise HTTPException(status_code=500, detail="Failed to extract segments")
    
@app.get("/export/{source_id}", response_class=FileResponse)
def export_translation(source_id: int):
    try:
        segments = get_latest_segments([source_id])
        if not segments:
            raise HTTPException(
                status_code=404, detail="No translated segments found.")

        # create a new document
        doc = Document()
        doc.add_heading("Translated Document", level=1)

        for segment in sorted(segments, key=lambda s: s['order']):
            doc.add_paragraph(segment["text"])

        # save the document to a temporary file
        file_path = f"/tmp/translated_{source_id}.docx"
        doc.save(file_path)

        return FileResponse(file_path, filename=f"translated_{source_id}.docx",
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Error generating DOCX: {str(e)}")


####### DICTIONARY, RULES, PROMPT.
@app.post("/prompt", response_model=str)
async def get_prompt(request: PromptRequest, dict = Depends(get_user_info)):
    try:
        if request.dictionary_id is not None:
            return build_prompt_from_dictionary(
                    request.dictionary_id,
                    request.dictionary_timestamp)
        else:
            # Return default task prompt - validate required fields
            required(request.original_language, "original_language is required when dictionary_id is not set.")
            required(request.translated_language, "translated_language is required when dictionary_id is not set.")
            return get_task_prompt(
                request.original_language,
                request.additional_sources_languages,
                request.translated_language)
    except HTTPException:
        raise
    except Exception as e:
        logger.error("Error getting prompt: %s", e)
        raise HTTPException(status_code=500, detail=f"Failed to get prompt: {str(e)}")

@app.get("/dictionaries", response_model=list[dict])
async def read_dictionaries(dictionary_id: int | None = None, dictionary_timestamp: int | None = None, user_info: dict = Depends(get_user_info)):
  return get_dictionaries(dictionary_id, dictionary_timestamp)

@app.post("/dictionaries", response_model=dict)
async def post_dictionary(dictionary: dict, user_info: dict = Depends(get_user_info)):
  username = user_info["preferred_username"]
  now = datetime.now(timezone.utc)
  if not "id" in dictionary or not dictionary["id"]:
    cursor = db.execute_sql("SELECT nextval('dictionary_id_seq')")
    dictionary_id = cursor.fetchone()[0]

    updated_dictionary = Dictionaries.create(
      id=dictionary_id,
      username=username,
      timestamp=now,
      **dictionary,
    ) 
  else:
    updated_dictionary = (Dictionaries
      .select()
      .where(Dictionaries.id == dictionary["id"])
      .order_by(Dictionaries.timestamp.desc())
      .limit(1)
      .get_or_none())
    if updated_dictionary is None:
      raise HTTPException(status_code=404, detail="Dictionary not found")
    updated_dictionary.username = username
    updated_dictionary.timestamp = now
    logger.info('d %s %s', dictionary, type(dictionary)) 
    logger.info('before %s %s', updated_dictionary, type(updated_dictionary)) 
    apply_dict(updated_dictionary, dictionary, logger)
    logger.info('after %s %s', updated_dictionary, type(updated_dictionary)) 
    updated_dictionary.save(force_insert=True)

  return get_dictionaries(updated_dictionary.id, epoch_microseconds(updated_dictionary.timestamp))[0]

@app.post("/dictionaries/prompt", response_model=dict)
async def create_prompt_dictionary(request: dict, user_info: dict = Depends(get_user_info)):
    try:
        name = request.get("name", None)
        original_language = request.get("original_language", None)
        additional_sources_languages = request.get("additional_sources_languages", [])
        translated_language = request.get("translated_language", None)

        required(original_language, "original_language is required.")
        required(translated_language, "translated_language is required.")

        timestamp=datetime.now(timezone.utc)

        cursor = db.execute_sql("SELECT nextval('dictionary_id_seq')")
        dictionary_id = cursor.fetchone()[0]

        created_dictionary = Dictionaries.create(
            id=dictionary_id,
            timestamp=timestamp,
            username=user_info['preferred_username'],
            name=name,
            original_language=original_language,
            additional_sources_languages=additional_sources_languages,
            translated_language=translated_language,
        )

        # Create default task prompt rule
        cursor = db.execute_sql("SELECT nextval('rule_id_seq')")
        rule_id = cursor.fetchone()[0]

        prompt_rule = Rules.create(
            id=rule_id,
            timestamp=timestamp,
            name="Default task prompt",
            username=user_info['preferred_username'],
            dictionary_id=dictionary_id,
            order=0,
            type=RULE_TYPE_TEXT,
            properties={"text": get_task_prompt(original_language, additional_sources_languages, translated_language)},
        )

        # Create paragraphs suffix rule
        cursor = db.execute_sql("SELECT nextval('rule_id_seq')")
        rule_id = cursor.fetchone()[0]

        paragraphs_rule = Rules.create(
            id=rule_id,
            timestamp=timestamp,
            name="Paragraphs suffix rule.",
            username=user_info['preferred_username'],
            dictionary_id=dictionary_id,
            order=1,
            type=RULE_TYPE_SEGMENTS_SUFFIX,
            properties={"text": SEGMENTS_SUFFIX},
        )

        return get_dictionaries(created_dictionary.id, epoch_microseconds(created_dictionary.timestamp))[0]
    except Exception as e:
        logger.error("Error adding or creating new dictionary: %s", e)
        raise HTTPException(status_code=500, detail=f"Failed adding or creating new dictionary: {str(e)}")


@app.get("/rules", response_model=list[dict])
def fetch_rules(dictionary_id: int | None = None, dictionary_timestamp: int | None = None, user_info: dict = Depends(get_user_info)):
  return get_rules(dictionary_id, dictionary_timestamp)


@app.post("/rules", response_model=list[dict])
async def post_rules(request: dict, user_info: dict = Depends(get_user_info)):
  username = user_info["preferred_username"]
  now = datetime.now(timezone.utc)
  rules = request.get("rules", [])

  if not rules:
    raise HTTPException(status_code=400, detail="No rules provided")

  updated_rule_ids = []
  for rule in rules:
    if not "id" in rule or not rule["id"]:
      # Create new rule
      cursor = db.execute_sql("SELECT nextval('rule_id_seq')")
      rule_id = cursor.fetchone()[0]

      updated_rule = Rules.create(
        id=rule_id,
        username=username,
        timestamp=now,
        **rule,
      )
    else:
      # Update existing rule
      updated_rule = (Rules
        .select()
        .where(Rules.id == rule["id"])
        .order_by(Rules.timestamp.desc())
        .limit(1)
        .get_or_none())
      if updated_rule is None:
        raise HTTPException(status_code=404, detail=f"Rule with id {rule['id']} not found")
      updated_rule.username = username
      updated_rule.timestamp = now
      apply_dict(updated_rule, rule)
      logger.info("RULE TIMESTAMP BEFORE SAVE [%d]", updated_rule.timestamp) 
      updated_rule.save(force_insert=True)

    updated_rule_ids.append(updated_rule.id)

  # Return all updated rules with full details
  return get_rules(rule_ids=updated_rule_ids)
